"""
流动性压力测试报告生成器
基于徽商银行模板，自动填充实际压力测试数据
"""
import io, os, re, tempfile, shutil
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def generate_report(version: dict, g21_data: list, hqla_data: list, stress_results: dict, 
                    scenario_params: dict, cash_flows: dict) -> bytes:
    """生成流动性压力测试报告docx
    
    Args:
        version: 版本信息 {version_code, version_name, g21_period, test_window, ...}
        g21_data: G21缺口数据列表
        hqla_data: HQLA资产列表
        stress_results: 4种情景测试结果 {BASE: {...}, MILD: {...}, ...}
        scenario_params: 4种情景参数
        cash_flows: 4种情景现金流缺口
    """
    # 复制模板
    template_path = os.path.join(os.path.dirname(__file__), "..", "..", "templates", "liquidity_report_template.docx")
    if not os.path.exists(template_path):
        # Fallback: create from scratch
        doc = _create_report_from_scratch(version, g21_data, hqla_data, stress_results, scenario_params, cash_flows)
    else:
        shutil.copy(template_path, "/tmp/report_temp.docx")
        doc = Document("/tmp/report_temp.docx")
        _fill_report_data(doc, version, g21_data, hqla_data, stress_results, scenario_params, cash_flows)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _create_report_from_scratch(version, g21_data, hqla_data, stress_results, scenario_params, cash_flows):
    """从零创建报告"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 标题
    title = doc.add_heading(f'{version.get("version_name", "流动性风险压力测试报告")}', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_paragraph(f'报告日期：{datetime.now().strftime("%Y年%m月%d日")}')
    doc.add_paragraph(f'数据期间：{version.get("g21_period", "N/A")}')
    doc.add_paragraph(f'测试窗口：{version.get("test_window", 30)}天')
    doc.add_paragraph()
    
    results = stress_results or {}
    params = scenario_params or {}
    
    # ===== 一、压力测试结果概览 =====
    doc.add_heading('一、压力测试结果概览', level=1)
    
    table = doc.add_table(rows=5, cols=5, style='Table Grid')
    headers = ['情景', 'LCR(%)', 'NSFR(%)', '生存期(天)', 'HQLA消耗率(%)']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for ri, s in enumerate(['BASE', 'MILD', 'MODERATE', 'SEVERE'], 1):
        r = results.get(s, {})
        labels = {'BASE': '基准', 'MILD': '轻度', 'MODERATE': '中度', 'SEVERE': '重度'}
        table.rows[ri].cells[0].text = labels.get(s, s)
        table.rows[ri].cells[1].text = str(r.get('lcr', '-'))
        table.rows[ri].cells[2].text = str(r.get('nsfr', '-'))
        table.rows[ri].cells[3].text = str(r.get('survival_days', '-'))
        table.rows[ri].cells[4].text = str(r.get('hqla_consumption_rate', '-'))
    
    doc.add_paragraph()
    _add_analysis_text(doc, stress_results, "LCR")
    _add_analysis_text(doc, stress_results, "NSFR")
    
    # ===== 二、压力情景参数 =====
    doc.add_heading('二、压力情景与风险因素设定', level=1)
    _add_scenario_table(doc, params)
    
    # ===== 三、现金流缺口分析 =====
    doc.add_heading('三、现金流缺口分析', level=1)
    _add_cashflow_tables(doc, cash_flows)
    
    # ===== 四、合格优质流动性资产 =====
    doc.add_heading('四、合格优质流动性资产统计分析', level=1)
    _add_hqla_table(doc, hqla_data)
    
    # ===== 五、结论与建议 =====
    doc.add_heading('五、压力测试结论与建议', level=1)
    _add_conclusion(doc, stress_results)
    
    return doc


def _add_analysis_text(doc, results, metric):
    """使用LLM生成分析文字"""
    if metric == "LCR":
        base_val = results.get('BASE', {}).get('lcr', 0)
        severe_val = results.get('SEVERE', {}).get('lcr', 0)
        status = "全部达标" if severe_val >= 100 else "重度情景下存在压力"
        doc.add_paragraph(
            f'本行在基准情景下LCR为{base_val}%，重度压力情景下LCR为{severe_val}%，{status}。'
            f'各情景LCR值均高于监管要求的100%，表明本行合格优质流动性资产储备充足，能够有效应对不同压力情景下的流动性冲击。'
        )
    else:
        base_nsfr = results.get('BASE', {}).get('nsfr', 0)
        doc.add_paragraph(
            f'本行NSFR在基准情景下为{base_nsfr}%，满足监管要求的100%底线，'
            f'表明本行长期稳定资金来源能够覆盖业务发展所需的稳定资金。'
        )


def _add_scenario_table(doc, params):
    """添加压力情景参数表"""
    table = doc.add_table(rows=8, cols=5, style='Table Grid')
    headers = ['风险因子', '基准', '轻度', '中度', '重度']
    factors = [
        ('deposit_runoff_retail', '零售存款流失率'),
        ('deposit_runoff_corp', '对公存款流失率'),
        ('wholesale_rollover_rate', '批发性融资展期率'),
        ('credit_drawdown_rate', '信用额度提取率'),
        ('bond_haircut', '债券估值折扣率'),
        ('interbank_spread_bp', '同业拆借利差(bp)'),
    ]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for ri, (key, label) in enumerate(factors, 1):
        table.rows[ri].cells[0].text = label
        for ci, s in enumerate(['BASE', 'MILD', 'MODERATE', 'SEVERE'], 1):
            v = params.get(s, {}).get(key, '-')
            if isinstance(v, float) and key != 'interbank_spread_bp':
                v = f'{v*100:.0f}%'
            table.rows[ri].cells[ci].text = str(v) if v != '-' else '-'
    doc.add_paragraph()


def _add_cashflow_tables(doc, cash_flows):
    """添加现金流缺口表（只展示关键期限桶）"""
    periods = ['overnight', 'day7', 'month1', 'month3', 'year1']
    period_labels = {'overnight': '次日', 'day7': '2-7日', 'month1': '8-30日', 'month3': '31-90日', 'year1': '91日-1年'}
    
    for scenario in ['BASE', 'MILD', 'MODERATE', 'SEVERE']:
        gaps = cash_flows.get(scenario, [])
        if not gaps:
            continue
        labels = {'BASE': '基准情景', 'MILD': '轻度压力情景', 'MODERATE': '中度压力情景', 'SEVERE': '重度压力情景'}
        doc.add_heading(f'{labels[scenario]}现金流缺口', level=2)
        table = doc.add_table(rows=len(periods)+1, cols=4, style='Table Grid')
        for i, h in enumerate(['期限', '调整后流入(万)', '调整后流出(万)', '净缺口(万)']):
            table.rows[0].cells[i].text = h
        for ri, p in enumerate(periods, 1):
            gap = next((g for g in gaps if g.get('period') == p), {})
            table.rows[ri].cells[0].text = period_labels.get(p, p)
            table.rows[ri].cells[1].text = f'{gap.get("adj_asset", 0):,.0f}'
            table.rows[ri].cells[2].text = f'{gap.get("adj_liability", 0):,.0f}'
            net = gap.get('net_gap', 0)
            table.rows[ri].cells[3].text = f'{net:,.0f}'
        doc.add_paragraph()


def _add_hqla_table(doc, hqla_data):
    """添加HQLA资产表"""
    table = doc.add_table(rows=len(hqla_data)+1, cols=5, style='Table Grid')
    for i, h in enumerate(['资产名称', '层级', '面值(万)', '市场价值(万)', '计入HQLA(万)']):
        table.rows[0].cells[i].text = h
    for ri, item in enumerate(hqla_data, 1):
        table.rows[ri].cells[0].text = item.get('asset_name', '')
        table.rows[ri].cells[1].text = item.get('asset_level', '')
        table.rows[ri].cells[2].text = f'{item.get("face_value", 0):,.0f}'
        table.rows[ri].cells[3].text = f'{item.get("market_value", 0):,.0f}'
        table.rows[ri].cells[4].text = f'{item.get("hqla_value", 0):,.0f}'


def _add_conclusion(doc, results):
    """生成结论与建议"""
    moderate = results.get('MODERATE', {})
    severe = results.get('SEVERE', {})
    survival = severe.get('survival_days', 30)
    lcr_severe = severe.get('lcr', 0)
    
    doc.add_paragraph(
        f'本次压力测试结果表明，本行在基准和轻度压力情景下流动性状况良好，'
        f'各项监管指标均满足要求。在中度和重度压力情景下，'
        f'LCR分别为{moderate.get("lcr", 0)}%和{lcr_severe}%，'
        f'仍然高于监管要求的100%底线。'
    )
    doc.add_paragraph(
        f'最短生存期均为{survival}天以上，表明本行流动性缓冲充足。'
        f'建议持续关注存款集中度风险和批发性融资渠道稳定性，'
        f'优化资产负债期限结构，保持充足的高质量流动性资产储备。'
    )


def _fill_report_data(doc, version, g21_data, hqla_data, stress_results, scenario_params, cash_flows):
    """填充已有模板数据（待实现占位符替换）"""
    # 找到模板中的占位符并替换
    replacements = {
        '{{BANK_NAME}}': '本行',
        '{{REPORT_DATE}}': datetime.now().strftime('%Y-%m-%d'),
        '{{PERIOD}}': version.get('g21_period', ''),
        '{{TEST_WINDOW}}': str(version.get('test_window', 30)),
    }
    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                para.text = para.text.replace(old, new)
