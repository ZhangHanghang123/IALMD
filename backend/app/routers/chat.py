"""智能对话分析 API — 接入 LLM (DeepSeek)，支持 SSE 流式响应"""
import json
import asyncio
import logging
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session
from sqlalchemy import desc
from ..database import get_db
from ..models.chat import IalmdChatSession, IalmdChatMessage
from ..models.system import SysLlmConfig
from ..schemas.common import ResponseBase
from ..dependencies import get_current_user
from ..services.llm_factory import get_llm

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/chat", tags=["智能对话"])

# ====================== Request Schema ======================

class SendMessageRequest(BaseModel):
    message: str = Field(..., min_length=1, max_length=8000, description="用户消息")

# ====================== System Prompt ======================

BANKING_SYSTEM_PROMPT = """你是一位资深的保险业经营分析专家，名叫「IALMD 分析师」。你的职责是帮助用户分析中国保险业的经营数据。

## 核心能力
1. **指标解读**：解读综合偿付能力充足率、核心偿付能力充足率、综合成本率(COR)、赔付率、退保率、新业务价值(NBV)、内含价值(EV)、总投资收益率、ROE等核心经营指标
2. **同业对比**：对比不同类型保险公司（保险集团、寿险、财险、再保险、健康险、养老险）的经营表现
3. **趋势分析**：分析保险公司指标的历史变化趋势和驱动因素
4. **归因分析**：解释指标变动的深层原因（监管政策、市场环境、保险公司策略等）

## 数据背景
- 覆盖中国主要保险公司（13家保险集团 + 75家寿险 + 88家财险 + 7家再保险 + 7家健康险 + 10家养老险）
- 数据来源：年报/半年报/季度报告/偿付能力报告/精算报告/保费收入公告/ESG报告等公开报告
- 时间范围：2016-2026年

## 图表输出（重要！）
当回复涉及数据排名、对比或多行数据时，使用以下图表标记语法来输出可视化图表：

### 柱状图（用于数据对比）
[chart:bar]
{"title":"头部险企2024年综合偿付能力充足率对比","items":[
  {"label":"新华保险","value":280},
  {"label":"中国人寿","value":262},
  {"label":"中国太保","value":248},
  {"label":"中国人保","value":235},
  {"label":"中国平安","value":210}
]}
[/chart]

### 排名图（用于并排排名）
[chart:rank]
{"title":"2024年上市险企总投资收益率TOP5","items":[
  {"name":"中国太平","value":5.8},
  {"name":"中国人寿","value":5.2},
  {"name":"中国太保","value":4.9},
  {"name":"中国平安","value":4.6},
  {"name":"新华保险","value":4.3}
]}
[/chart]

### 折线图（用于趋势分析）
[chart:line]
{"title":"中国人寿综合偿付能力充足率趋势","data":[
  {"label":"2020","value":275},
  {"label":"2021","value":268},
  {"label":"2022","value":260},
  {"label":"2023","value":258},
  {"label":"2024","value":262}
]}
[/chart]

每当你在回复中涉及排名/对比/趋势数据时，请积极使用以上图表语法来输出可视化内容。图表标记放在文字分析之后。

## 参考资料输出（重要！）
在回复的**最末尾**，你必须使用以下标记输出本次分析所参考的全部文档和数据来源：

[references]
{"items":[
  {"type":"report","title":"中国人寿2024年年度报告","bank":"中国人寿","year":"2024","reportType":"年度报告","detail":"保费收入、偿付能力充足率等核心指标数据来源"},
  {"type":"indicator","title":"综合偿付能力充足率","bank":"中国人寿","year":"2024","value":"262%","category":"偿付能力","detail":"衡量保险公司资本总体充足状况的核心监管指标"},
  {"type":"indicator","title":"综合成本率(COR)","bank":"人保财险","year":"2024","value":"96.5%","category":"业务质量","detail":"反映财险承保盈利能力的关键指标"},
  {"type":"external","title":"国家金融监督管理总局保险业偿付能力状况通报","source":"国家金融监督管理总局","year":"2024","detail":"行业整体偿付能力、风险综合评级等宏观参考数据"},
  {"type":"external","title":"中国保险行业协会经营数据","source":"中国保险行业协会","year":"2024","detail":"行业保费收入、赔付支出等统计数据"}
]}
[/references]

### 参考项类型说明
- **type="report"**: 参考的保险公开报告（年报/半年报/季度报告/偿付能力报告/精算报告/保费收入公告/ESG报告等），需填写bank/year/reportType
- **type="indicator"**: 参考的具体经营指标数据，需填写bank/year/value/category（偿付能力/业务质量/盈利能力/规模/价值/渠道/ESG）
- **type="external"**: 参考的外部数据来源（监管统计/行业报告/宏观政策文件等），需填写source/year

### 要求
- **尽可能详细**地列出所有参考的数据来源，包括涉及的每家保险公司的报告、每个引用的指标、每个外部数据源
- 同一保险公司的不同报告分别列出（如年报+偿付能力报告+精算报告）
- 同一指标的不同年份分别列出
- 参考资料标记放在所有文字和图表之后，作为回复的最后一部分

## 回复规范
- 用简洁专业的语言回答，适当使用 Markdown 表格和列表
- 涉及数据时注明数据来源和年份
- 对比分析时给出排序和关键差异，同时用图表增强展示
- 无法回答时如实说明，不要编造数据
- 对于预测性问题，基于公开信息和合理推演给出观点，并标注「仅供参考」

## 语气
- 专业但不生硬，像一位有经验的保险业分析师在和你讨论
- 使用中文，适当使用保险行业术语
"""

# ====================== 会话管理 ======================

@router.get("/sessions", response_model=ResponseBase)
def list_sessions(
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user),
):
    """获取用户对话会话列表"""
    sessions = (
        db.query(IalmdChatSession)
        .filter(
            IalmdChatSession.user_id == current_user.get("id"),
            IalmdChatSession.status == 1,
            IalmdChatSession.is_deleted == 0,
        )
        .order_by(desc(IalmdChatSession.updated_at))
        .limit(50)
        .all()
    )

    return ResponseBase(data=[
        {
            "id": s.id,
            "title": s.session_title,
            "type": s.session_type,
            "message_count": s.message_count,
            "updated_at": s.updated_at.isoformat() if s.updated_at else None,
            "created_at": s.created_at.isoformat() if s.created_at else None,
        }
        for s in sessions
    ])


@router.post("/sessions", response_model=ResponseBase)
def create_session(
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user),
):
    """创建新对话"""
    session = IalmdChatSession(
        user_id=current_user.get("id"),
        session_title="新对话",
        session_type="ANALYSIS",
        created_by=current_user.get("id"),
    )
    db.add(session)
    db.commit()
    db.refresh(session)

    return ResponseBase(data={
        "id": session.id,
        "title": session.session_title,
    })


@router.get("/messages/{session_id}", response_model=ResponseBase)
def get_messages(
    session_id: int,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user),
):
    """获取对话消息历史"""
    session = db.query(IalmdChatSession).filter(
        IalmdChatSession.id == session_id,
        IalmdChatSession.user_id == current_user.get("id"),
        IalmdChatSession.status == 1,
    ).first()

    if not session:
        raise HTTPException(status_code=404, detail="会话不存在")

    messages = (
        db.query(IalmdChatMessage)
        .filter(
            IalmdChatMessage.session_id == session_id,
            IalmdChatMessage.status == 1,
            IalmdChatMessage.is_deleted == 0,
        )
        .order_by(IalmdChatMessage.created_at.asc())
        .all()
    )

    return ResponseBase(data=[
        {
            "id": m.id,
            "role": m.role,
            "content": m.content,
            "message_type": m.message_type,
            "chart_json": m.chart_json,
            "table_json": m.table_json,
            "tokens_used": m.tokens_used,
            "model_name": m.model_name,
            "created_at": m.created_at.isoformat() if m.created_at else None,
        }
        for m in messages
    ])


@router.delete("/sessions/{session_id}", response_model=ResponseBase)
def delete_session(
    session_id: int,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user),
):
    """删除对话会话"""
    session = db.query(IalmdChatSession).filter(
        IalmdChatSession.id == session_id,
        IalmdChatSession.user_id == current_user.get("id"),
    ).first()

    if not session:
        raise HTTPException(status_code=404, detail="会话不存在")

    session.is_deleted = 1
    db.commit()
    return ResponseBase(message="删除成功")


# ====================== 智能对话（核心） ======================

# 发送消息路由 (使用 /send 后缀避免与 get_messages 冲突)
@router.get("/messages/{session_id}/send")
@router.post("/messages/{session_id}/send")
async def send_message(
    session_id: int,
    message: str | None = None,  # GET query param
    req_data: SendMessageRequest | None = None,  # POST body
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user),
):
    """发送消息并获取 AI 流式响应 (SSE)
    
    GET: /api/chat/messages/{session_id}/send?message=xxx (用于 EventSource)
    POST: /api/chat/messages/{session_id}/send with body {"message": "xxx"}
    """
    # 从 GET query 或 POST body 获取消息
    user_msg = None
    if message:
        user_msg = message
    elif req_data and req_data.message:
        user_msg = req_data.message
    
    if not user_msg:
        raise HTTPException(status_code=400, detail="请提供 message 参数")
    
    return await handle_chat_message(session_id, user_msg, db, current_user)


async def handle_chat_message(
    session_id: int,
    user_message: str,
    db: Session,
    current_user: dict,
):
    """处理聊天消息的内部函数（GET 和 POST 共用）"""
    
    user_message = user_message.strip()
    if not user_message:
        raise HTTPException(status_code=400, detail="消息不能为空")

    # 验证会话
    session = db.query(IalmdChatSession).filter(
        IalmdChatSession.id == session_id,
        IalmdChatSession.user_id == current_user.get("id"),
        IalmdChatSession.status == 1,
    ).first()

    if not session:
        raise HTTPException(status_code=404, detail="会话不存在")

    # 保存用户消息
    user_msg = IalmdChatMessage(
        session_id=session_id,
        role="USER",
        content=user_message,
        message_type="TEXT",
        created_by=current_user.get("id"),
    )
    db.add(user_msg)
    session.message_count = (session.message_count or 0) + 1

    # 首次对话自动生成标题
    if session.session_title == "新对话":
        title = user_message[:30]
        if len(user_message) > 30:
            title += "…"
        session.session_title = title

    db.commit()

    # 构建消息历史上下文（最近 20 条）
    history = (
        db.query(IalmdChatMessage)
        .filter(
            IalmdChatMessage.session_id == session_id,
            IalmdChatMessage.status == 1,
            IalmdChatMessage.is_deleted == 0,
        )
        .order_by(IalmdChatMessage.created_at.asc())
        .limit(40)  # 最近 20 轮对话
        .all()
    )

    # 获取 LLM 配置名
    llm_config_name = "DeepSeek"
    try:
        llm_cfg = db.query(SysLlmConfig).filter(SysLlmConfig.is_enabled == 1, SysLlmConfig.is_default == 1).first()
        if llm_cfg:
            llm_config_name = llm_cfg.model_name or llm_cfg.provider_name
    except Exception:
        pass

    # 构建消息列表
    from langchain_core.messages import SystemMessage, HumanMessage, AIMessage
    messages = [SystemMessage(content=BANKING_SYSTEM_PROMPT)]

    for h in history:
        if h.role == "USER":
            messages.append(HumanMessage(content=h.content))
        elif h.role == "ASSISTANT":
            messages.append(AIMessage(content=h.content))

    async def event_generator():
        full_response = ""
        tokens_used = 0
        error_occurred = False

        try:
            llm = get_llm()

            # 检查是否为 mock 模式
            from ..services.llm_mock import MockChatModel
            is_mock = isinstance(llm, MockChatModel)

            if is_mock:
                # Mock 模式：返回预设引导
                mock_text = (
                    "您好！我是 IALMD 分析师，目前平台运行在模拟模式。\n\n"
                    "模拟模式下我无法提供实时数据分析，但您可以使用以下功能：\n\n"
                    "**完全可用的功能**：\n"
                    "- 📊 工作流编排 — 在「工作流编排」页面创建和执行分析工作流\n"
                    "- 📋 指标库管理 — 在「经营指标库」页面查看和管理指标体系\n"
                    "- 🏦 银行信息管理 — 在「保险机构管理」页面浏览 47 家银行基础数据\n"
                    "- ⚙️ LLM 配置 — 在「系统设置」页面配置 DeepSeek API Key\n\n"
                    "**配置真实 API Key 后**，我将能够为您提供：\n"
                    "- 保险经营指标的深度解读\n"
                    "- 跨银行同业对比分析\n"
                    "- 历史趋势和归因分析\n"
                    "- 个性化经营诊断建议\n\n"
                    "请前往「系统设置 → LLM配置」启用 DeepSeek 并填入 API Key。"
                )
                full_response = mock_text
                tokens_used = 0

                # 逐段输出模拟流式效果
                chunk_size = 20
                for i in range(0, len(mock_text), chunk_size):
                    chunk = mock_text[i:i + chunk_size]
                    yield f"data: {json.dumps({'type': 'text', 'content': chunk, 'index': i // chunk_size}, ensure_ascii=False)}\n\n"
                    await asyncio.sleep(0.02)

            else:
                # 真实 LLM 流式调用
                stream = llm.astream(messages)
                chunk_index = 0
                async for chunk in stream:
                    if hasattr(chunk, "content") and chunk.content:
                        delta = chunk.content
                        full_response += delta
                        chunk_index += 1

                        # 获取 token 用量（如果有）
                        if hasattr(chunk, "response_metadata"):
                            usage = chunk.response_metadata.get("token_usage", {})
                            tokens_used = usage.get("total_tokens", 0)

                        yield f"data: {json.dumps({'type': 'text', 'content': delta, 'index': chunk_index}, ensure_ascii=False)}\n\n"

        except Exception as e:
            error_occurred = True
            logger.error(f"LLM 调用失败: {e}")
            error_msg = f"抱歉，分析服务暂时不可用：{str(e)[:200]}"
            full_response = error_msg
            yield f"data: {json.dumps({'type': 'error', 'content': error_msg}, ensure_ascii=False)}\n\n"

        # 保存 AI 回复到数据库
        if not error_occurred or full_response:
            try:
                ai_msg = IalmdChatMessage(
                    session_id=session_id,
                    role="ASSISTANT",
                    content=full_response,
                    message_type="TEXT",
                    tokens_used=tokens_used,
                    model_name=llm_config_name,
                    created_by=current_user.get("id"),
                )
                db.add(ai_msg)
                session.message_count = (session.message_count or 0) + 1
                session.updated_at = datetime.now()
                db.commit()
            except Exception as e:
                logger.error(f"保存 AI 回复失败: {e}")

        # 发送结束信号
        yield f"data: {json.dumps({'type': 'done', 'session_title': session.session_title, 'tokens_used': tokens_used}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
            "Access-Control-Allow-Origin": "*",
        },
    )


@router.get("/sessions/{session_id}/export")
def export_session(session_id: int, db: Session = Depends(get_db),
                   current_user: dict = Depends(get_current_user),
                   message_id: int | None = None):
    """导出对话为Word文档（message_id指定时仅导出该条）"""
    from docx import Document
    from docx.shared import Pt
    import io
    from fastapi.responses import StreamingResponse as SR

    session = db.query(IalmdChatSession).filter(
        IalmdChatSession.id == session_id, IalmdChatSession.is_deleted == 0
    ).first()
    if not session:
        return ResponseBase(code=404, message="会话不存在")

    msg_query = db.query(IalmdChatMessage.content, IalmdChatMessage.role,
                        IalmdChatMessage.created_at).filter(
        IalmdChatMessage.session_id == session_id, IalmdChatMessage.is_deleted == 0
    )
    if message_id:
        msg_query = msg_query.filter(IalmdChatMessage.id == message_id)
    messages = msg_query.order_by(IalmdChatMessage.created_at).all()

    doc = Document()
    doc.styles['Normal'].font.size = Pt(11)
    doc.add_heading(session.session_title or '保险经营分析报告', level=0)
    doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}')

    import re, json
    for content, role, created_at in messages:
        label = '用户' if role == 'USER' else 'IALMD分析师'
        doc.add_heading(f'{label}', level=2)
        
        # 解析内容：文本段落 / 表格 / 图表
        _render_content_to_docx(doc, content)

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    from urllib.parse import quote
    fn = quote(f"智能对话分析报告_{session.session_title or 'export'}.docx")
    return SR(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
              headers={"Content-Disposition": f"attachment; filename={fn}"})


# ====================== 参考资料导出与下载 ======================

# 机构名称映射（简称 → 服务器目录名）
_BANK_NAME_MAP = {
    "工商银行": "中国工商银行", "工行": "中国工商银行",
    "建设银行": "中国建设银行", "建行": "中国建设银行",
    "农业银行": "中国农业银行", "农行": "中国农业银行",
    "中国银行": "中国银行", "中行": "中国银行",
    "交通银行": "交通银行", "交行": "交通银行",
    "邮储银行": "中国邮政储蓄银行", "邮政储蓄银行": "中国邮政储蓄银行",
    "招商银行": "招商银行", "招行": "招商银行",
    "中信银行": "中信银行",
    "光大银行": "中国光大银行",
    "民生银行": "中国民生银行",
    "浦发银行": "上海浦东发展银行", "上海浦东发展银行": "上海浦东发展银行",
    "上海银行": "上海银行",
    "兴业银行": "兴业银行",
    "平安银行": "平安银行",
    "华夏银行": "华夏银行",
    "北京银行": "北京银行",
    "兰州银行": "兰州银行",
}

# 报告类型映射
_REPORT_TYPE_MAP = {
    "年度报告": "年度报告", "年报": "年度报告",
    "半年度报告": "半年度报告", "半年报": "半年度报告",
    "季度报告": "季度报告", "季报": "季度报告", "一季报": "季度报告", "三季报": "季度报告",
    "业绩快报": "业绩快报",
    "资本充足率报告": "资本充足率信息披露报告", "资本充足率信息披露报告": "资本充足率信息披露报告",
    "流动性报告": "流动性风险信息披露报告", "流动性风险信息披露报告": "流动性风险信息披露报告",
    "ESG报告": "社会责任报告ESG", "社会责任报告": "社会责任报告ESG", "社会责任报告ESG": "社会责任报告ESG",
    "普惠金融报告": "普惠金融服务报告", "普惠金融服务报告": "普惠金融服务报告",
    "绿色金融报告": "绿色金融专项报告", "绿色金融专项报告": "绿色金融专项报告",
    "消费者权益保护报告": "消费者权益保护工作报告", "消费者权益保护工作报告": "消费者权益保护工作报告",
}

import os
DATA_DIR = os.environ.get("IALMD_DATA_DIR", "/opt/ialmd/data")
REPORT_DIR = os.path.join(DATA_DIR, "保险经营报告下载")


class ReferenceExportRequest(BaseModel):
    """参考资料导出请求"""
    items: list = Field(..., description="参考资料列表")
    session_title: str = Field(default="智能对话分析", description="对话标题")


@router.post("/references/export")
def export_references(
    req: ReferenceExportRequest,
    current_user: dict = Depends(get_current_user),
):
    """将参考资料导出为 Excel 文件"""
    import io as _io
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from fastapi.responses import StreamingResponse as SR
    from urllib.parse import quote

    wb = Workbook()

    # 样式定义
    header_font = Font(name="Microsoft YaHei", size=11, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="1677FF", end_color="1677FF", fill_type="solid")
    cell_font = Font(name="Microsoft YaHei", size=10)
    thin_border = Border(
        left=Side(style="thin", color="D9D9D9"),
        right=Side(style="thin", color="D9D9D9"),
        top=Side(style="thin", color="D9D9D9"),
        bottom=Side(style="thin", color="D9D9D9"),
    )
    wrap_align = Alignment(wrap_text=True, vertical="top")

    type_labels = {"report": "银行报告", "indicator": "经营指标", "external": "外部数据源"}
    type_colors = {"report": "52C41A", "indicator": "1677FF", "external": "FA8C16"}

    # Sheet 1: 参考资料总览
    ws1 = wb.active
    ws1.title = "参考资料总览"
    headers1 = ["序号", "类型", "标题", "银行", "年份", "数值", "分类/来源", "说明"]
    for ci, h in enumerate(headers1, 1):
        c = ws1.cell(row=1, column=ci, value=h)
        c.font = header_font
        c.fill = header_fill
        c.alignment = Alignment(horizontal="center", vertical="center")

    for ri, item in enumerate(req.items, 2):
        t = item.get("type", "")
        row_data = [
            ri - 1,
            type_labels.get(t, t),
            item.get("title", ""),
            item.get("bank", ""),
            item.get("year", ""),
            item.get("value", ""),
            item.get("category", "") or item.get("source", "") or item.get("reportType", ""),
            item.get("detail", ""),
        ]
        for ci, val in enumerate(row_data, 1):
            c = ws1.cell(row=ri, column=ci, value=val)
            c.font = cell_font
            c.border = thin_border
            c.alignment = wrap_align
        # 类型颜色标记
        color = type_colors.get(t, "8C8C8C")
        ws1.cell(row=ri, column=2).fill = PatternFill(start_color=color, end_color=color, fill_type="solid")
        ws1.cell(row=ri, column=2).font = Font(name="Microsoft YaHei", size=10, bold=True, color="FFFFFF")

    # 列宽
    for col, width in zip("ABCDEFGH", [6, 12, 35, 14, 8, 12, 18, 40]):
        ws1.column_dimensions[col].width = width

    # Sheet 2: 参考报告详情
    reports = [it for it in req.items if it.get("type") == "report"]
    if reports:
        ws2 = wb.create_sheet("参考报告")
        headers2 = ["序号", "报告标题", "银行", "报告类型", "年份", "说明"]
        for ci, h in enumerate(headers2, 1):
            c = ws2.cell(row=1, column=ci, value=h)
            c.font = header_font
            c.fill = header_fill
            c.alignment = Alignment(horizontal="center", vertical="center")
        for ri, item in enumerate(reports, 2):
            for ci, val in enumerate([
                ri - 1, item.get("title", ""), item.get("bank", ""),
                item.get("reportType", ""), item.get("year", ""), item.get("detail", ""),
            ], 1):
                c = ws2.cell(row=ri, column=ci, value=val)
                c.font = cell_font
                c.border = thin_border
                c.alignment = wrap_align
        for col, width in zip("ABCDEF", [6, 35, 14, 16, 8, 40]):
            ws2.column_dimensions[col].width = width

    # Sheet 3: 参考指标详情
    indicators = [it for it in req.items if it.get("type") == "indicator"]
    if indicators:
        ws3 = wb.create_sheet("参考指标")
        headers3 = ["序号", "指标名称", "银行", "年份", "数值", "指标分类", "说明"]
        for ci, h in enumerate(headers3, 1):
            c = ws3.cell(row=1, column=ci, value=h)
            c.font = header_font
            c.fill = header_fill
            c.alignment = Alignment(horizontal="center", vertical="center")
        for ri, item in enumerate(indicators, 2):
            for ci, val in enumerate([
                ri - 1, item.get("title", ""), item.get("bank", ""),
                item.get("year", ""), item.get("value", ""),
                item.get("category", ""), item.get("detail", ""),
            ], 1):
                c = ws3.cell(row=ri, column=ci, value=val)
                c.font = cell_font
                c.border = thin_border
                c.alignment = wrap_align
        for col, width in zip("ABCDEFG", [6, 20, 14, 8, 12, 14, 40]):
            ws3.column_dimensions[col].width = width

    # Sheet 4: 外部数据源
    externals = [it for it in req.items if it.get("type") == "external"]
    if externals:
        ws4 = wb.create_sheet("外部数据源")
        headers4 = ["序号", "数据源名称", "来源机构", "年份", "说明"]
        for ci, h in enumerate(headers4, 1):
            c = ws4.cell(row=1, column=ci, value=h)
            c.font = header_font
            c.fill = header_fill
            c.alignment = Alignment(horizontal="center", vertical="center")
        for ri, item in enumerate(externals, 2):
            for ci, val in enumerate([
                ri - 1, item.get("title", ""), item.get("source", ""),
                item.get("year", ""), item.get("detail", ""),
            ], 1):
                c = ws4.cell(row=ri, column=ci, value=val)
                c.font = cell_font
                c.border = thin_border
                c.alignment = wrap_align
        for col, width in zip("ABCDE", [6, 35, 20, 8, 40]):
            ws4.column_dimensions[col].width = width

    buf = _io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    fn = quote(f"参考资料_{req.session_title}.xlsx")
    return SR(buf, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
              headers={"Content-Disposition": f"attachment; filename={fn}"})


@router.get("/references/report-download")
def download_report(
    bank: str,
    report_type: str,
    year: str,
    current_user: dict = Depends(get_current_user),
):
    """下载银行报告 PDF 文件"""
    import glob as _glob
    from fastapi.responses import FileResponse, StreamingResponse as SR
    import io as _io

    # 映射机构名称
    dir_name = _BANK_NAME_MAP.get(bank, bank)
    # 映射报告类型
    rpt_type = _REPORT_TYPE_MAP.get(report_type, report_type)

    search_dir = os.path.join(REPORT_DIR, dir_name, rpt_type)
    if not os.path.isdir(search_dir):
        raise HTTPException(status_code=404, detail=f"未找到报告目录: {dir_name}/{rpt_type}")

    # 搜索匹配年份的 PDF
    patterns = [
        os.path.join(search_dir, f"*{year}*年度报告*.pdf"),
        os.path.join(search_dir, f"*{year}*.pdf"),
        os.path.join(search_dir, f"{year}*.pdf"),
    ]
    found = None
    for pat in patterns:
        matches = _glob.glob(pat)
        if matches:
            found = matches[0]
            break

    if not found:
        # 列出可用年份
        all_pdfs = sorted(_glob.glob(os.path.join(search_dir, "*.pdf")))
        available = [os.path.basename(f) for f in all_pdfs[:10]]
        raise HTTPException(
            status_code=404,
            detail=f"未找到 {dir_name} {year}年{rpt_type}。可用文件: {', '.join(available)}"
        )

    from urllib.parse import quote
    basename = os.path.basename(found)
    fn = quote(basename)
    return FileResponse(
        found,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={fn}"}
    )


@router.get("/references/report-check")
def check_report(
    bank: str,
    report_type: str,
    year: str,
    current_user: dict = Depends(get_current_user),
):
    """检查银行报告文件是否可下载"""
    import glob as _glob
    dir_name = _BANK_NAME_MAP.get(bank, bank)
    rpt_type = _REPORT_TYPE_MAP.get(report_type, report_type)
    search_dir = os.path.join(REPORT_DIR, dir_name, rpt_type)

    if not os.path.isdir(search_dir):
        return ResponseBase(data={"available": False, "reason": "目录不存在"})

    patterns = [
        os.path.join(search_dir, f"*{year}*.pdf"),
    ]
    for pat in patterns:
        matches = _glob.glob(pat)
        if matches:
            return ResponseBase(data={"available": True, "file": os.path.basename(matches[0])})

    # 返回可用年份
    all_pdfs = sorted(_glob.glob(os.path.join(search_dir, "*.pdf")))
    years_avail = []
    for f in all_pdfs:
        bn = os.path.basename(f)
        for y in ["2024", "2023", "2022", "2021", "2020", "2019", "2018", "2017", "2016", "2015"]:
            if y in bn:
                years_avail.append(y)
                break
    return ResponseBase(data={"available": False, "available_years": years_avail[:10]})


def _render_content_to_docx(doc, content: str):
    """将聊天内容渲染到docx：文本段落/HTML表格/chart图表/references参考资料"""
    import re, json
    from docx.shared import Pt, RGBColor

    # 提取并移除 [references]...[/references] 结构（避免原始标记出现在正文）
    ref_pattern = re.compile(r'\[references\]([\s\S]*?)\[/references\]', re.DOTALL)
    ref_items = None
    m = ref_pattern.search(content)
    if m:
        content = content[:m.start()] + content[m.end():]
        try:
            ref_data = json.loads(m.group(1).strip())
            ref_items = ref_data.get("items", [])
        except Exception:
            ref_items = None

    segments = []
    last_idx = 0
    chart_pattern = re.compile(r'\[chart:(\w+)\]([\s\S]*?)\[/chart\]', re.DOTALL)
    table_pattern = re.compile(r'\|(?:.+\|)+\s*\n\|[-:\s|]+\|\s*\n(?:\|.+\|\s*\n?)+')
    
    for m in chart_pattern.finditer(content):
        if m.start() > last_idx:
            segments.append(('text', content[last_idx:m.start()]))
        try:
            chart_type = m.group(1)  # bar / rank / line
            chart_json = json.loads(m.group(2).strip())
            chart_json['_chart_type'] = chart_type
            segments.append(('chart', chart_json))
        except:
            segments.append(('text', m.group(0)))
        last_idx = m.end()
    if last_idx < len(content):
        segments.append(('text', content[last_idx:]))
    
    for seg_type, seg_data in segments:
        if seg_type == 'text':
            # 按表格分割文本段
            text_parts = table_pattern.split(seg_data)
            tables_in_text = table_pattern.findall(seg_data)
            
            for i, tp in enumerate(text_parts):
                _add_text_to_docx(doc, tp)
                if i < len(tables_in_text):
                    _add_md_table_to_docx(doc, tables_in_text[i])
        elif seg_type == 'chart':
            _add_chart_image_to_docx(doc, seg_data)

    # 渲染参考资料章节
    if ref_items:
        _add_references_to_docx(doc, ref_items)


def _add_references_to_docx(doc, ref_items: list):
    """将参考资料渲染为结构化的 Word 章节"""
    from docx.shared import Pt, RGBColor

    if not ref_items:
        return

    doc.add_heading('参考资料', level=2)

    type_label = {"report": "📄 报告", "indicator": "📊 指标", "external": "🌐 外部数据"}
    type_order = ["report", "indicator", "external"]

    for t in type_order:
        items = [x for x in ref_items if x.get("type") == t]
        if not items:
            continue
        doc.add_heading(type_label.get(t, t), level=3)
        for item in items:
            title = item.get("title", "")
            # 构建元信息行
            meta_parts = []
            if item.get("bank"):
                meta_parts.append(item["bank"])
            if item.get("year"):
                meta_parts.append(str(item["year"]))
            if item.get("reportType"):
                meta_parts.append(item["reportType"])
            if item.get("value"):
                meta_parts.append(item["value"])
            if item.get("category"):
                meta_parts.append(item["category"])
            if item.get("source"):
                meta_parts.append(item["source"])

            # 标题行（加粗）
            p = doc.add_paragraph()
            run = p.add_run(f"• {title}")
            run.bold = True
            run.font.size = Pt(10.5)

            # 元信息行（灰色）
            if meta_parts:
                p2 = doc.add_paragraph()
                run2 = p2.add_run("    " + " | ".join(meta_parts))
                run2.font.size = Pt(9)
                run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            # 详情行
            if item.get("detail"):
                p3 = doc.add_paragraph()
                run3 = p3.add_run("    " + item["detail"])
                run3.font.size = Pt(9)
                run3.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


def _add_text_to_docx(doc, text: str):
    """添加文本段落"""
    for line in text.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith('#'):
            level = min(len(stripped) - len(stripped.lstrip('#')), 3)
            doc.add_heading(stripped.lstrip('#').strip(), level=level)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped.startswith('> '):
            p = doc.add_paragraph(stripped[2:])
            p.runs[0].italic = True
        else:
            doc.add_paragraph(stripped)


def _add_md_table_to_docx(doc, md_table: str):
    """将Markdown表格转换为docx表格"""
    from docx.shared import Pt
    lines = md_table.strip().split('\n')
    if len(lines) < 3:
        return
    headers = [c.strip().replace('**', '') for c in lines[0].split('|') if c.strip()]
    rows_data = []
    for line in lines[2:]:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows_data.append(cells)
    
    if not headers or not rows_data:
        return
    
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers), style='Table Grid')

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
    
    # 数据行
    for ri, row in enumerate(rows_data):
        for ci, val in enumerate(row):
            if ci < len(headers):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = val
                for p in cell.paragraphs:
                    p.runs[0].font.size = Pt(9)
    
    doc.add_paragraph()  # 表格后空行


def _add_chart_image_to_docx(doc, chart_data: dict):
    """使用 matplotlib 生成图表图片并插入 Word 文档"""
    import matplotlib
    matplotlib.use('Agg')  # 无头模式，不弹窗
    import matplotlib.pyplot as plt
    from docx.shared import Inches
    import io as _io

    # ---- 配置中文字体 ----
    _setup_matplotlib_fonts()

    title = chart_data.get('title', '')
    items = chart_data.get('items', []) or chart_data.get('data', [])
    if not items or len(items) < 2:
        return

    chart_type = chart_data.get('_chart_type', 'bar')
    names = [it.get('name', it.get('label', str(i))) for i, it in enumerate(items)]
    values = [it.get('value', 0) for it in items]

    # 颜色方案
    palette = ['#1677ff', '#52c41a', '#fa8c16', '#722ed1', '#eb2f96',
               '#13c2c2', '#f5222d', '#2f54eb', '#a0d911', '#faad14']

    fig, ax = plt.subplots(figsize=(7.5, 4.2), dpi=150)
    fig.patch.set_facecolor('white')

    if chart_type == 'line':
        # ---- 折线图 ----
        ax.plot(range(len(names)), values, color='#1677ff', linewidth=2,
                marker='o', markersize=6, markerfacecolor='#1677ff',
                markeredgecolor='white', markeredgewidth=1.5)
        ax.fill_between(range(len(names)), values, alpha=0.12, color='#1677ff')
        ax.set_xticks(range(len(names)))
        ax.set_xticklabels(names, fontsize=9)
        for i, v in enumerate(values):
            ax.annotate(f'{v:,.2f}', (i, v), textcoords="offset points",
                        xytext=(0, 10), ha='center', fontsize=8, color='#333')
        ax.set_ylim(bottom=min(values) * 0.9 if min(values) > 0 else min(values) * 1.1,
                    top=max(values) * 1.15)
        ax.grid(axis='y', linestyle='--', alpha=0.3)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

    elif chart_type == 'rank':
        # ---- 排名图（水平条形图，降序）----
        paired = sorted(zip(values, names), key=lambda x: x[0], reverse=True)
        sorted_vals = [p[0] for p in paired]
        sorted_names = [p[1] for p in paired]
        colors = [palette[i % len(palette)] for i in range(len(paired))]
        bars = ax.barh(range(len(sorted_names)), sorted_vals, color=colors,
                       height=0.6, edgecolor='white', linewidth=0.5)
        ax.set_yticks(range(len(sorted_names)))
        ax.set_yticklabels(sorted_names, fontsize=9)
        ax.invert_yaxis()
        for i, (bar, v) in enumerate(zip(bars, sorted_vals)):
            ax.text(v + max(sorted_vals) * 0.01, bar.get_y() + bar.get_height() / 2,
                    f'{v:,.2f}', va='center', ha='left', fontsize=8, color='#333')
        ax.set_xlim(right=max(sorted_vals) * 1.15)
        ax.grid(axis='x', linestyle='--', alpha=0.3)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

    else:
        # ---- 柱状图（默认）----
        x_pos = range(len(names))
        colors = [palette[i % len(palette)] for i in range(len(names))]
        bars = ax.bar(x_pos, values, color=colors, width=0.6,
                      edgecolor='white', linewidth=0.5)
        ax.set_xticks(x_pos)
        ax.set_xticklabels(names, fontsize=9, rotation=30, ha='right')
        for bar, v in zip(bars, values):
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + max(values) * 0.01,
                    f'{v:,.2f}', ha='center', va='bottom', fontsize=8, color='#333')
        ax.set_ylim(top=max(values) * 1.15)
        ax.grid(axis='y', linestyle='--', alpha=0.3)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

    ax.set_title(title, fontsize=12, fontweight='bold', pad=12, color='#1a1a1a')
    plt.tight_layout()

    buf = _io.BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight', dpi=150,
                facecolor='white', edgecolor='none')
    plt.close(fig)
    buf.seek(0)

    doc.add_picture(buf, width=Inches(5.8))
    doc.add_paragraph()  # 图后空行


def _setup_matplotlib_fonts():
    """配置 matplotlib 中文字体（只需执行一次）"""
    import matplotlib
    if getattr(matplotlib, '_ialmd_font_configured', False):
        return
    import matplotlib.font_manager as fm
    import os, glob

    # 候选中文字体路径（按优先级）
    candidates = []
    # Linux 常见中文字体
    for pat in ['/usr/share/fonts/**/*wqy*', '/usr/share/fonts/**/*Noto*CJK*',
                '/usr/share/fonts/**/*SimHei*', '/usr/share/fonts/**/*Droid*Fallback*',
                '/usr/share/fonts/**/*msyh*', '/usr/share/fonts/truetype/**/*CJK*']:
        candidates.extend(glob.glob(pat, recursive=True))
    # Windows 常见中文字体
    win_fonts = [r'C:\Windows\Fonts\msyh.ttc', r'C:\Windows\Fonts\simhei.ttf',
                 r'C:\Windows\Fonts\simsun.ttc']
    candidates.extend([f for f in win_fonts if os.path.exists(f)])

    font_name = None
    for fp in candidates:
        if os.path.exists(fp):
            try:
                fm.fontManager.addfont(fp)
                prop = fm.FontProperties(fname=fp)
                font_name = prop.get_name()
                break
            except Exception:
                continue

    if font_name:
        matplotlib.rcParams['font.sans-serif'] = [font_name]
    else:
        # 没有中文字体时用默认（中文可能显示为方框，但不影响图表结构）
        matplotlib.rcParams['font.sans-serif'] = ['DejaVu Sans', 'sans-serif']

    matplotlib.rcParams['axes.unicode_minus'] = False
    matplotlib._ialmd_font_configured = True
